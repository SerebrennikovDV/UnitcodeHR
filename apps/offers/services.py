"""Сервисы офферов: генерация .docx-документа оффера."""
import io
from datetime import date

from django.core.files.base import ContentFile
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_offer_docx(offer):
    """Формирует .docx с офертой и привязывает к offer.document."""
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)

    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('ПРЕДЛОЖЕНИЕ О РАБОТЕ (JOB OFFER)')
    r.bold = True
    r.font.size = Pt(14)

    doc.add_paragraph(f'Дата: {date.today().strftime("%d.%m.%Y")}')
    doc.add_paragraph()

    candidate = offer.application.candidate
    vacancy = offer.application.vacancy

    doc.add_paragraph(f'Уважаемый(ая) {candidate.first_name} {candidate.last_name}!')
    doc.add_paragraph()
    doc.add_paragraph(
        f'ООО «Юниткод» рад предложить вам работу на позиции '
        f'«{vacancy.title}».'
    )

    doc.add_paragraph('Условия предложения:')
    bullets = [
        f'Должность: {vacancy.title}',
        f'Заработная плата: {offer.salary:,.0f} руб.',
        f'Планируемая дата выхода: {offer.start_date.strftime("%d.%m.%Y")}',
        f'Испытательный срок: {offer.probation_months} мес.',
        'Формат работы: удалённый',
    ]
    for b in bullets:
        doc.add_paragraph(b, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph(
        'Для подтверждения принятия предложения, пожалуйста, ответьте на '
        'это письмо в течение 3 рабочих дней. С условиями договора '
        '(трудовой / гражданско-правовой / самозанятого) вас ознакомит '
        'HR-менеджер.'
    )
    doc.add_paragraph()
    doc.add_paragraph('С уважением,')
    doc.add_paragraph('Команда ООО «Юниткод»')

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    filename = f'offer_{offer.pk}_{candidate.last_name}.docx'
    offer.document.save(filename, ContentFile(buf.read()), save=True)
    return offer.document.name
