"""Аудит финального отчёта на риск детекции ИИ-генерации.

Подсчёт типичных «маркеров» ИИ-текста: канцеляризмы, безличные обороты,
повторы клише, длинные сложные предложения. Не заменяет AI-детектор,
но показывает, какие абзацы переписать в первую очередь.

Запуск:
    .\\venv\\Scripts\\python.exe scripts\\ai_risk_audit.py
"""
import sys, io, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

REPORT = r'C:\Users\admin\Desktop\Учеба\Серебренников_ПД_Отчет.docx'

# Канцеляризмы и клише, по которым ChatGPT/Claude легко детектятся
AI_MARKERS = [
    'в данном разделе', 'данный раздел', 'в данном пункте',
    'следует отметить, что', 'необходимо отметить',
    'таким образом, можно сделать вывод', 'таким образом,',
    'в рамках данной', 'в рамках практики',
    'в качестве',
    'позволяет обеспечить', 'позволяет достичь',
    'играет ключевую роль', 'играет важную роль',
    'представляет собой', 'является важным',
    'в современных условиях', 'в современном мире',
    'на сегодняшний день',
    'обеспечивает возможность',
    'осуществляется', 'осуществляет',
    'реализован функционал',
    'учитывая вышеизложенное',
    'необходимо учитывать',
    'способствует',
    'является одним из',
    'на основе анализа',
    'в ходе работы', 'в ходе практики',
    'результаты показывают',
    'данный подход',
    'эффективное решение',
    'оптимальное решение',
    'комплексный подход',
    'современные технологии',
    'актуальной задачей',
]


def audit():
    doc = Document(REPORT)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    table_text = []
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        table_text.append(p.text.strip())

    all_chunks = paragraphs + table_text
    total_chars = sum(len(p) for p in all_chunks)
    total_words = sum(len(p.split()) for p in all_chunks)

    print(f'=== ИТОГИ ОТЧЁТА ===')
    print(f'Параграфов с текстом: {len(paragraphs)}')
    print(f'Текстовых ячеек таблиц: {len(table_text)}')
    print(f'Всего символов: {total_chars:,}')
    print(f'Всего слов: {total_words:,}')
    print()

    # Глобальный счёт маркеров
    total_text = ' '.join(all_chunks).lower()
    marker_counts = {}
    for m in AI_MARKERS:
        c = total_text.count(m)
        if c:
            marker_counts[m] = c
    print('=== ТОП ИИ-маркеров (по всему отчёту) ===')
    for m, c in sorted(marker_counts.items(), key=lambda x: -x[1])[:25]:
        print(f'  {c:3} раз — "{m}"')

    # Топ-15 параграфов с наибольшей плотностью маркеров на 100 слов
    risky = []
    for i, p in enumerate(paragraphs):
        pl = p.lower()
        words = max(len(p.split()), 1)
        score = sum(pl.count(m) for m in AI_MARKERS)
        density = score / words * 100
        if score >= 2 and words >= 25:
            risky.append((density, score, words, i, p))
    risky.sort(reverse=True)

    print('\n=== ТОП-15 АБЗАЦЕВ ДЛЯ ПЕРЕПИСКИ (по плотности ИИ-клише) ===')
    for rank, (dens, sc, words, idx, p) in enumerate(risky[:15], 1):
        print(f'\n#{rank}  плотность={dens:.1f} клише/100слов  маркеров={sc}  слов={words}  P{idx}')
        snippet = p[:350] + ('…' if len(p) > 350 else '')
        print(f'   "{snippet}"')


if __name__ == '__main__':
    audit()
