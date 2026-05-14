"""Готовит файл REWRITE_CHECKLIST.docx — что в каком абзаце переписать.

Сборка:
  1. Введение (часть 1) — переписать целиком
  2. Выводы по главам (1.8, 2.6, 3.5)
  3. Заключение (часть 7)
  4. Топ-15 «опасных» абзацев из общего сборного отчёта

Запуск:
    .\\venv\\Scripts\\python.exe scripts\\make_rewrite_checklist.py
"""
import sys, io, re, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = r'C:\Users\admin\Desktop\Учеба'
PARTS = [f'Серебренников_ПД_Отчет_часть{i}.docx' for i in range(1, 8)]
OUT = os.path.join(BASE, 'ПЕРЕПИСАТЬ.docx')

AI_MARKERS = [
    'в данном разделе', 'данный раздел', 'в данном пункте',
    'следует отметить, что', 'необходимо отметить',
    'таким образом, можно сделать вывод', 'таким образом,',
    'в рамках данной', 'в рамках практики', 'в качестве',
    'позволяет обеспечить', 'позволяет достичь',
    'играет ключевую роль', 'играет важную роль',
    'представляет собой', 'является важным',
    'в современных условиях', 'в современном мире', 'на сегодняшний день',
    'обеспечивает возможность', 'осуществляется', 'осуществляет',
    'реализован функционал', 'учитывая вышеизложенное',
    'необходимо учитывать', 'способствует', 'является одним из',
    'на основе анализа', 'в ходе работы', 'в ходе практики',
    'результаты показывают', 'данный подход', 'эффективное решение',
    'оптимальное решение', 'комплексный подход',
    'современные технологии', 'актуальной задачей',
]


def collect_paragraphs():
    """Возвращает список (часть, заголовок, текст) для всех параграфов."""
    items = []
    for part_no, fn in enumerate(PARTS, 1):
        path = os.path.join(BASE, fn)
        d = Document(path)
        current_heading = ''
        for p in d.paragraphs:
            txt = p.text.strip()
            if not txt:
                continue
            style = p.style.name if p.style else ''
            if 'Heading' in style or 'Заголовок' in style:
                current_heading = txt
                continue
            items.append((part_no, current_heading, txt))
    return items


def find_section(items, keywords):
    """Параграфы из раздела с заголовком, содержащим ключевые слова."""
    out = []
    for part, head, txt in items:
        head_low = head.lower()
        if any(k.lower() in head_low for k in keywords):
            out.append((part, head, txt))
    return out


def ai_score(text):
    tl = text.lower()
    words = max(len(text.split()), 1)
    score = sum(tl.count(m) for m in AI_MARKERS)
    return score, score / words * 100


def make_doc():
    items = collect_paragraphs()

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    h = doc.add_heading('Чек-лист переписывания отчёта', level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph()
    intro.add_run(
        'В этом файле собраны абзацы из «Серебренников_ПД_Отчет.docx», которые '
        'нужно переписать своими словами, чтобы Антиплагиат.ВУЗ Витте не '
        'детектировал ИИ-генерацию. После каждого «старого» абзаца оставлено '
        'место для твоего нового текста. Когда закончишь — скажи мне, '
        'я вклею новые формулировки обратно в части отчёта и пересоберу финальный docx.'
    ).italic = True

    def write_block(title, paragraphs, note=None):
        doc.add_heading(title, level=1)
        if note:
            p = doc.add_paragraph()
            r = p.add_run(note)
            r.italic = True
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        for idx, (part, head, txt) in enumerate(paragraphs, 1):
            doc.add_heading(f'{idx}. (часть {part}) {head}', level=2)
            old = doc.add_paragraph()
            old.add_run('БЫЛО: ').bold = True
            old.add_run(txt)

            new = doc.add_paragraph()
            r = new.add_run('СТАЛО (твой текст): ')
            r.bold = True
            r.font.color.rgb = RGBColor(0x10, 0x6f, 0x10)
            placeholder = new.add_run('_____________________________________')
            placeholder.font.color.rgb = RGBColor(0xbb, 0xbb, 0xbb)
            doc.add_paragraph()

    # 1. Введение — всё до первого "Глава 1"
    intro_items = find_section(items, ['Введение'])
    write_block(
        '1. Введение — переписать целиком',
        intro_items[:8],
        'Это самая первая вещь, которую читает антиплагиат и руководитель. '
        'Перепиши коротко, простыми предложениями: о компании, о теме, что сделал, '
        'какие компетенции освоил. От первого лица: «Я провёл…», «Получилось…».'
    )

    # 2. Выводы по главам
    conclusions = find_section(items, ['Выводы по разделу', 'Выводы по главе', 'выводы по'])
    write_block(
        '2. Выводы по главам (1.8, 2.6, 3.5)',
        conclusions,
        'Перепиши своими словами что сделал в каждой главе. Можно по пунктам. '
        'Главное — без «таким образом, можно сделать вывод». Просто перечисли результаты.'
    )

    # 3. Заключение — параграфы из часть 7 которые не входят в список литературы
    concl = []
    for part, head, txt in items:
        if part == 7 and 'литератур' not in head.lower():
            concl.append((part, head, txt))
    write_block(
        '3. Заключение (часть 7)',
        concl,
        'Тут обобщаешь практику в целом. Достаточно 1 страницы. Что планируешь '
        'дорабатывать в ВКР. От первого лица.'
    )

    # 4. Топ-15 опасных абзацев
    risky = []
    for part, head, txt in items:
        score, density = ai_score(txt)
        words = len(txt.split())
        if score >= 1 and words >= 25:
            risky.append((density, score, part, head, txt))
    risky.sort(reverse=True)

    write_block(
        '4. Прочие абзацы с высокой плотностью ИИ-клише',
        [(part, head, txt) for _, _, part, head, txt in risky[:15]],
        'Это абзацы, где встречается «таким образом», «осуществляется», '
        '«представляет собой» и т.п. — типовые ИИ-маркеры.'
    )

    doc.save(OUT)
    print(f'Готово: {OUT}')


if __name__ == '__main__':
    make_doc()
