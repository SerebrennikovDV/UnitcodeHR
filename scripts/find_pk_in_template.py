"""Поиск таблиц компетенций в шаблоне отчёта."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
d = Document(r'C:\Users\admin\Desktop\Учеба\ФИО_Отчет_БИ ЦЭ_new.docx')

print('=== Paragraphs containing ПК or компетенц ===')
for i, p in enumerate(d.paragraphs):
    if 'ПК' in p.text or 'компетенц' in p.text.lower():
        print(f'P{i}: {p.text[:300]}')

print('\n=== Tables with PK-like content ===')
for ti, t in enumerate(d.tables):
    table_text = ''
    for row in t.rows:
        for cell in row.cells:
            table_text += cell.text + ' '
    if 'ПК' in table_text or 'компетенц' in table_text.lower():
        print(f'\n--- TABLE {ti} (rows={len(t.rows)}) ---')
        for ri, row in enumerate(t.rows):
            cells = [c.text.strip().replace('\n', ' | ')[:250] for c in row.cells]
            print(f'  R{ri}: {cells}')
