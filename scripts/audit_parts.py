"""Аудит частей отчёта на соответствие требованиям ВУЗа."""
import sys, io, os, re
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

base = r'C:\Users\admin\Desktop\Учеба'
parts = sorted([f for f in os.listdir(base)
                if f.startswith('Серебренников_ПД_Отчет_часть') and f.endswith('.docx')])

for fn in parts:
    path = os.path.join(base, fn)
    d = Document(path)
    text = ' '.join(p.text for p in d.paragraphs)
    for t in d.tables:
        for row in t.rows:
            for cell in row.cells:
                text += ' ' + cell.text
    pks = re.findall(r'ПК-?\d+', text)
    pk_counts = {}
    for pk in pks:
        n = re.sub(r'ПК-?', 'ПК-', pk)
        pk_counts[n] = pk_counts.get(n, 0) + 1
    gost = 'ГОСТ 34.602' in text
    bpmn_as = 'AS IS' in text or 'AS-IS' in text or 'КАК ЕСТЬ' in text
    bpmn_to = 'TO BE' in text or 'TO-BE' in text or 'КАК ДОЛЖНО' in text
    er = 'ER-диаграмм' in text or 'ER диаграмм' in text or 'сущность-связь' in text.lower()
    nfbk = 'НФБК' in text or 'нормальной форме' in text.lower() or 'нормализаци' in text.lower()
    test_plan = 'тест-план' in text.lower() or 'тест план' in text.lower()
    test_case = 'тест-кейс' in text.lower() or 'тестовый сценарий' in text.lower()
    bug = 'баг-репорт' in text.lower() or 'дефект' in text.lower()
    print(f'=== {fn} ===')
    print(f'  paragraphs={len(d.paragraphs)}  tables={len(d.tables)}  chars={len(text)}')
    print(f'  ПК: {sorted(pk_counts.items())}')
    flags = []
    if gost: flags.append('GOST34.602')
    if bpmn_as: flags.append('AS-IS')
    if bpmn_to: flags.append('TO-BE')
    if er: flags.append('ER')
    if nfbk: flags.append('NF/normalization')
    if test_plan: flags.append('test-plan')
    if test_case: flags.append('test-case')
    if bug: flags.append('bug-report')
    print(f'  flags: {flags}')
